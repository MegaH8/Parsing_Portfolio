from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
from lxml import html
from bs4 import BeautifulSoup
import time
from random import randrange

headers = {
    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/106.0.0.0 Safari/537.36'
}


def tables(items, power, name):
    # создание пустого документа
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    paragraph = doc.add_paragraph(name)
    paragraph.bold = True
    paragraph.alignment = 1

    table = doc.add_table(1, 2)

    # определяем стиль таблицы
    table.style = 'Table Grid'
    # Получаем строку с колонками из добавленной таблицы
    head_cells = table.rows[0].cells
    col = table.columns[0]
    col.width = Cm(9.43)
    col = table.columns[1]
    col.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    col.width = Cm(5.75)
    for i, item in enumerate(['Мощность основной источник (PRIME)', power]):
        p = head_cells[i].paragraphs[0]
        p.add_run(item)
    # добавляем данные к существующей таблице
    for row in items:
        # убирает пустую строку
        if None in row:
            continue
        # добавляем строку с ячейками к объекту таблицы
        cells = table.add_row().cells
        for i, item in enumerate(row):
            # добавляем строку с ячейками к объекту таблицы

            # вставляем данные в ячейки
            cells[i].text = str(item)
            # если последняя ячейка
            if i == 1:
                cells[i].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save('test.docx')


def get_articles_urls(url):
    s = requests.Session()
    response = s.get(url=url,headers=headers)

    soup = BeautifulSoup(response.text, 'lxml')
    pagination_count = int(soup.find('div', class_='navigation-pages').find_all('a')[-1].text.strip())

    articles_urls_list = []
    for page in range(1, pagination_count + 1):
    #for page in range(1, 2):
        response = s.get(url=f'https://1kwt.com/katalog/dizelnye-generatory/?PAGEN_1={page}', headers=headers)
        soup = BeautifulSoup(response.text, 'lxml')

        articles_url = soup.find_all('div', class_='img-product')

        for au in articles_url:
            art_url = au.a.get('href')
            articles_urls_list.append(f'https://1kwt.com{art_url}')

        time.sleep(randrange(2,5))
        print(f'обработал {page}/{pagination_count}')

    with open('articles_urls.txt', 'w') as file:
        for url in articles_urls_list:
            file.write(f'{url}\n')

    return 'Работа по сбору ссылок выполнена'


def get_data(file_path):
    with open(file_path) as file:
        urls_list = [line.strip() for line in file.readlines()]

    s = requests.Session()

    for url in urls_list[:3]:
        response = s.get(url=url, headers=headers)
        soup = BeautifulSoup(response.text, 'lxml')

        prod_name = soup.find(class_='bx-title').text
        prod_power = soup.find_all(class_='har_main_hr')[1].find(class_='val-props').text + ' кВт'
        coef_pw = soup.find_all('div', class_='har_main_hr')[11].find(class_='val-props').text
        frequency = soup.find_all('div', class_='har_main_hr')[12].find(class_='val-props').text + ' Гц'
        voltage = soup.find_all('div', class_='har_main_hr')[2].find(class_='val-props').text
        phases = soup.find_all('div', class_='har_main_hr')[7].find(class_='val-props').text
        current = soup.find_all('div', class_='har_main_hr')[13].find(class_='val-props').text + ' А'



        # данные таблицы
        items = (
            ('Коэффициент мощности', coef_pw),
            ('Частота', frequency),
            ('Напряжение', voltage),
            ('Количество фаз', phases),
            ('Номинальный ток', current),
            ('Производитель двигателя', '1288'),
            ('Модель двигателя', '1288'),
            ('Кол-во расположение цилиндров', '1288'),
            ('Номинальная мощность', '1288'),
            ('Объём двигателя', '1288'),
            ('Частота вращения', '1288'),
            ('Тип охлаждения', '1288'),
            ('Объём системы охлаждения', '1288'),
            ('Объем масляной системы', 'Объем масляной системы'),
            ('Расход топлива при 50/75/100 % мощности', '1288'),
            ('Объем топливного бака', '1288'),
            ('Производитель альтернатора', '1288'),
            ('Модель альтернатора', '1288'),
            ('Номинальная мощность', '1288'),
            ('Тип альтернатора', '1288'),
            ('Класс изоляции', '1288'),
            ('Степень защиты', '1288'),
            ('Габаритные размеры (открытое исполнение)', '1288'),
            ('Габаритные размеры (открытое исполнение)', '1288'),
            ('Вес (открытое исполнение)', '1288')
        )
        tables(items, prod_power, prod_name)


def main():
    #print(get_articles_urls(url='https://1kwt.com/katalog/dizelnye-generatory/?PAGEN_1=1'))

    get_data('articles_urls.txt')


if __name__ == '__main__':
    main()
